import numpy as np
import matplotlib.pyplot as plt
import math
import random

from docx import Document
from docx.shared import Inches


def cell_dist(cell1, cell2):
    '''
    Funkcja oblicza odległość między wybranymi członkami układu.
    cell1: pierwszy wybrany członek układu
    cell2: drugi wybrany członek układu
    Zwraca obliczoną odległość.
    '''
    # współrzędne w wierszu i kolumnie
    row1, col1 = cell1
    row2, col2 = cell2
    
    row_dist = abs(row2 - row1)
    col_dist = abs(col2 - col1)
    
    dist = row_dist**2 + col_dist**2

    return math.sqrt(dist)
    
    
def plot_table_color(color_arr, color_names, t):
    '''
    Funkcja rysuje tablicę opinii poszczególnych członków układu.
    color_arr: tablica opinii poszczególnych członków
    color_names: nazwy kolorów (konwersja nazw na zrozumiałą dla matplotlib)
    t: krok symulacji
    '''
    fig, ax = plt.subplots()

    for i in range(color_arr.shape[0]):
        for j in range(color_arr.shape[1]):
            ax.add_patch(plt.Rectangle((j, color_arr.shape[0]-i-1), 1, 1, color=color_names[color_arr[i][j]][0]))

    ax.set_aspect('equal')
    ax.set_xlim(0, color_arr.shape[1])
    ax.set_ylim(0, color_arr.shape[0])

    ax.set_xticks([])
    ax.set_yticks([])

    plt.savefig('./wyniki/3/tabela_zad3.png')
    
    doc.add_paragraph(f'\nt = {t:}')
    doc.add_picture('./wyniki/3/tabela_zad3.png', width=Inches(4))
    doc.save('./wyniki/3/plot_document_zad3.docx')


# stworzenie nowego dokumentu
doc = Document()

# możliwe opinie do przyjęcia przez członków grupy (układu)
# K = 3
colors = ['R', 'G', 'B']
color_names = {
    'R': 'red',
    'G': 'green',
    'B': 'blue'
}

# wpływy poszczególnych opinii na każdego członka grupy (układu)
idx_opinion = {i: [] for i in range(1, 401)}

# rozmiar układu
L = 20
arr_dim = L * L
# informacje o przyjętej opinii oraz oddziaływań z członkami o tych samych i przeciwnych opiniach
idx_color_p_s = {}
for i in range(1, arr_dim + 1):
    s = float(format(random.uniform(0, 1), '.3f'))
    p = float(format(1 - s, '.3f'))
    opinion = random.choice(colors)
    idx_color_p_s.update({i: [opinion, s, p]})

# układ
values = np.arange(1, 401)
model_arr = values.reshape(20, 20)

# początkowy krok układu
t = 0

alpha = 3

color_arr_rep = []
color_arr = np.array([[idx_color_p_s[num][0] for num in row] for row in model_arr])
# dodaj do sprawdzenia obecny stan układu
color_arr_rep.append(color_arr)

plot_table_color(color_arr, color_names, t)

while True:
    for i in range(model_arr.shape[0]):
        for j in range(model_arr.shape[1]):
            # nr komórki i jej współrzędne
            curr_cell = model_arr[i][j]
            curr_cell_coords = (i, j)
            
            # oblicz siłę oddziaływań od poszczególnych opiniic
            opinion_red = 4 * sum([(idx_color_p_s[model_arr[i][j]][1]) / (1 + cell_dist(curr_cell_coords, (i, j))**alpha) 
                                if idx_color_p_s[model_arr[i][j]][0] == idx_color_p_s[curr_cell][0]   # jeśli kolor innej opinii jest taki sam
                                else (idx_color_p_s[model_arr[i][j]][2]) / (1 + cell_dist(curr_cell_coords, (i, j))**alpha)  # jeśli kolor innej opinii jest inny
                            for i in range(model_arr.shape[0])
                            for j in range(model_arr.shape[1])
                                if idx_color_p_s[model_arr[i][j]][0] == 'R'])
            
            opinion_green = 4 * sum([(idx_color_p_s[model_arr[i][j]][1]) / (1 + cell_dist(curr_cell_coords, (i, j))**alpha) 
                                if idx_color_p_s[model_arr[i][j]][0] == idx_color_p_s[curr_cell][0]
                                else (idx_color_p_s[model_arr[i][j]][2]) / (1 + cell_dist(curr_cell_coords, (i, j))**alpha)
                            for i in range(model_arr.shape[0])
                            for j in range(model_arr.shape[1])
                                if idx_color_p_s[model_arr[i][j]][0] == 'G'])
            
            opinion_blue = 4 * sum([(idx_color_p_s[model_arr[i][j]][1]) / (1 + cell_dist(curr_cell_coords, (i, j))**alpha) 
                                if idx_color_p_s[model_arr[i][j]][0] == idx_color_p_s[curr_cell][0]
                                else (idx_color_p_s[model_arr[i][j]][2]) / (1 + cell_dist(curr_cell_coords, (i, j))**alpha)
                            for i in range(model_arr.shape[0])
                            for j in range(model_arr.shape[1])
                                if idx_color_p_s[model_arr[i][j]][0] == 'B'])

            opinion_red = round(opinion_red, 2)
            opinion_green = round(opinion_green, 2)
            opinion_blue = round(opinion_blue, 2)

            # wpływy wywierane na członka przez poszczególne opinie..
            cell_opinion_lst = [opinion_red, opinion_green, opinion_blue]
            # dodaj je do danych o opiniach aktualnego członka + wskaż jego nową opinię (największy wpływ)
            idx_opinion[curr_cell] = [cell_opinion_lst, colors[cell_opinion_lst.index(max(cell_opinion_lst))]]

    # zmnień kolor na najmocniej oddziaływujący na poszczególnych członków
    for i in range(model_arr.shape[0]):
        for j in range(model_arr.shape[1]):
            curr_cell = model_arr[i][j]
            idx_color_p_s[curr_cell][0] = idx_opinion[curr_cell][1]

    color_arr = np.array([[idx_color_p_s[num][0] for num in row] for row in model_arr])
    
    t += 1
    
    color_arr_rep.append(color_arr)
    
    # zakończ gdy nie ma zmian opinii/ występuje migotanie
    if len(color_arr_rep) == 4:
        if np.equal(color_arr_rep[0], color_arr_rep[1]).all():
            break
        # migotwanie
        elif np.equal(color_arr_rep[0], color_arr_rep[2]).all() and np.equal(color_arr_rep[1], color_arr_rep[3]).all():
            print('migotanie')
            break
        else:
            color_arr_rep.clear()  
            
    if color_arr_rep:
        plot_table_color(color_arr_rep[-1], color_names, t)
    else:
        plot_table_color(color_arr, color_names, t)
