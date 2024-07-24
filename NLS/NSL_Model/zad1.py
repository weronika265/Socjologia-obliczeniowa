import numpy as np
import matplotlib.pyplot as plt
import math

from docx import Document
from docx.shared import Inches


def cell_dist(cell1, cell2):
    '''
    Funkcja oblicza odległość między wybranymi członkami układu.
    cell1: pierwszy wybrany członek układu
    cell2: drugi wybrany członek układu
    Zwraca obliczoną odległość.
    '''
    # współrzędne w wierszu i kolumnie
    row1, col1 = cell1
    row2, col2 = cell2
    
    row_dist = abs(row2 - row1)
    col_dist = abs(col2 - col1)
    
    dist = row_dist**2 + col_dist**2

    return math.sqrt(dist)


def plot_table(model_arr, idx_color, color_names, idx_opinion, t):
    '''
    Funkcja rysuje tablicę opinii poszczególnych członków układu.
    model_arr: tablica wartości poszczególnych członków
    idx_color: opinie poszczególnych członków
    color_names: nazwy kolorów (konwersja nazw na zrozumiałą dla matplotlib)
    idx_opinion: wartości opinii dla członka układu pochodzące od każdej opinii
    t: krok symulacji
    '''
    fig, ax = plt.subplots()

    for i in range(model_arr.shape[0]):
        for j in range(model_arr.shape[1]):
            ax.add_patch(plt.Rectangle((j, model_arr.shape[0]-i-1), 1, 1, color=color_names[idx_color[model_arr[i][j]]]))
            ax.text(j + 0.5, model_arr.shape[0]-i-0.5, f'{model_arr[i][j]}', color='black',
                    ha='center', va='center', fontsize=12)

    ax.set_aspect('equal')
    ax.set_xlim(0, model_arr.shape[1])
    ax.set_ylim(0, model_arr.shape[0])

    ax.set_xticks([])
    ax.set_yticks([])

    plt.savefig('./wyniki/1/tabela_zad1.png')
    
    doc.add_paragraph(f'\nt = {t:}')
    
    if t != 'Init':
        for i, (lst, opinion) in idx_opinion.items():
            doc.add_paragraph(f'Członek nr {i} - wpływy [R, G, B]: {lst}, wybrana opinia: {opinion}')
    
    doc.add_picture('./wyniki/1/tabela_zad1.png', width=Inches(4))
    doc.save('./wyniki/1/plot_document_zad1.docx')
 
 
# stworzenie nowego dokumentu
doc = Document()

# możliwe opinie do przyjęcia przez członków grupy (układu)
# K = 3
colors = ['R', 'G', 'B']
color_names = {
    'R': 'red',
    'G': 'green',
    'B': 'blue'
}

# opinie przyjęte przez członków grupy (układu)
idx_color = {
    1: 'R',
    2: 'G',
    3: 'R',
    4: 'R',
    5: 'B',
    6: 'B',
    7: 'R',
    8: 'G',
    9: 'B',
}

# wpływy poszczególnych opinii na każdego członka grupy (układu)
idx_opinion = {
    1: [],
    2: [],
    3: [],
    4: [],
    5: [],
    6: [],
    7: [],
    8: [],
    9: [],
}

# układ
model_arr = np.array([[1, 2, 3], [4, 5, 6], [7, 8, 9]])

# początkowy krok układu
t = 0

# macierz z opiniami członków ukłądu
color_arr = np.array([[idx_color[num] for num in row] for row in model_arr])

plot_table(model_arr, idx_color, color_names, idx_opinion, t = 'Init')

while True:
    color_arr_prev = color_arr.copy()
    for i in range(model_arr.shape[0]):
        for j in range(model_arr.shape[1]):
            # nr komórki i jej współrzędne
            curr_cell = model_arr[i][j]
            curr_cell_coords = (i, j)
            
            # oblicz siłę oddziaływań od poszczególnych opinii
            opinion_red = 4 * sum([(model_arr[i][j] / 10) / (1 + cell_dist(curr_cell_coords, (i, j))**2) 
                                if idx_color[model_arr[i][j]] == idx_color[curr_cell]   # jeśli kolor innej opinii jest taki sam
                                else (1 - model_arr[i][j] / 10) / (1 + cell_dist(curr_cell_coords, (i, j))**2)  # jeśli kolor innej opinii jest inny
                            for i in range(model_arr.shape[0])
                            for j in range(model_arr.shape[1])
                                if idx_color[model_arr[i][j]] == 'R'])

            opinion_green = 4 * sum([(model_arr[i][j] / 10) / (1 + cell_dist(curr_cell_coords, (i, j))**2)
                                    if idx_color[model_arr[i][j]] == idx_color[curr_cell]
                                    else (1 - model_arr[i][j] / 10) / (1 + cell_dist(curr_cell_coords, (i, j))**2)
                            for i in range(model_arr.shape[0])
                            for j in range(model_arr.shape[1])
                                if idx_color[model_arr[i][j]] == 'G'])

            opinion_blue = 4 * sum([(model_arr[i][j] / 10) / (1 + cell_dist(curr_cell_coords, (i, j))**2)
                                    if idx_color[model_arr[i][j]] == idx_color[curr_cell]
                                    else (1 - model_arr[i][j] / 10) / (1 + cell_dist(curr_cell_coords, (i, j))**2)
                            for i in range(model_arr.shape[0])
                            for j in range(model_arr.shape[1])
                                if idx_color[model_arr[i][j]] == 'B'])

            opinion_red = round(opinion_red, 2)
            opinion_green = round(opinion_green, 2)
            opinion_blue = round(opinion_blue, 2)

            # wpływy wywierane na członka przez poszczególne opinie..
            cell_opinion_lst = [opinion_red, opinion_green, opinion_blue]
            # dodaj je do danych o opiniach aktualnego członka + wskaż jego nową opinię (największy wpływ)
            idx_opinion[curr_cell] = [cell_opinion_lst, colors[cell_opinion_lst.index(max(cell_opinion_lst))]]

    # zmnień kolor na najmocniej oddziaływujący na poszczególnych członków
    for i in range(model_arr.shape[0]):
        for j in range(model_arr.shape[1]):
            curr_cell = model_arr[i][j]
            idx_color[curr_cell] = idx_opinion[curr_cell][1]

    color_arr = np.array([[idx_color[num] for num in row] for row in model_arr])
    
    plot_table(model_arr, idx_color, color_names, idx_opinion, t)
    
    t += 1
    
    if np.equal(color_arr, color_arr_prev).all():
        break
